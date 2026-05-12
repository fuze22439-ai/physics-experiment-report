"""
通用 Word 模板填充脚本
支持两种模式：
1. 简单文本替换（替换段落中的占位符）
2. docxtpl Jinja2 渲染

用法:
    python fill_template.py 模板.docx 输出.docx --data '{"key":"value"}'
    python fill_template.py 模板.docx 输出.docx --replace '{"占位符":"替换文本"}'
"""
import json
import sys
import os


def fill_by_replace(template_path, output_path, replacements: dict):
    """简单文本替换：将段落中的占位符替换为指定文本"""
    from docx import Document
    doc = Document(template_path)

    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                # 尝试在 run 级别替换以保留格式
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        break
                else:
                    # fallback：段落级别
                    para.text = para.text.replace(old, new)

    doc.save(output_path)
    return True


def fill_by_jinja(template_path, output_path, context: dict):
    """Jinja2 模板渲染（需要模板中有 {{ }} 标签）"""
    from docxtpl import DocxTemplate
    doc = DocxTemplate(template_path)
    doc.render(context)
    doc.save(output_path)
    return True


def auto_fill(template_path, output_path, data: dict):
    """自动检测模板类型并填充"""
    from docx import Document
    doc = Document(template_path)
    full_text = " ".join(p.text for p in doc.paragraphs)

    if "{{" in full_text or "{%" in full_text:
        # Jinja2 模板
        print("  检测到 Jinja2 模板，使用 docxtpl 渲染")
        return fill_by_jinja(template_path, output_path, data)
    else:
        # 简单替换
        print("  使用文本替换模式")
        return fill_by_replace(template_path, output_path, data)


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='Word 模板填充工具')
    parser.add_argument('template', help='模板 .docx 路径')
    parser.add_argument('output', help='输出 .docx 路径')
    parser.add_argument('--data', type=str, default='{}',
                        help='JSON 数据，如 \'{"title":"扫描光电流","name":"张三"}\'')
    parser.add_argument('--replace', type=str, default=None,
                        help='替换映射 JSON，如 \'{"占位符":"替换内容"}\'')
    parser.add_argument('--jinja', action='store_true',
                        help='强制使用 Jinja2 模式')

    args = parser.parse_args()

    if not os.path.exists(args.template):
        print(f"错误：模板不存在 {args.template}")
        sys.exit(1)

    try:
        data = json.loads(args.data)
    except json.JSONDecodeError:
        print("错误：--data JSON 格式无效")
        sys.exit(1)

    if args.replace:
        try:
            replacements = json.loads(args.replace)
        except json.JSONDecodeError:
            print("错误：--replace JSON 格式无效")
            sys.exit(1)

        print(f"填充模板：{args.template}")
        print(f"替换映射：{replacements}")
        fill_by_replace(args.template, args.output, replacements)
    elif args.jinja:
        print(f"Jinja2 渲染：{args.template}")
        print(f"数据：{data}")
        fill_by_jinja(args.template, args.output, data)
    else:
        print(f"自动填充：{args.template}")
        auto_fill(args.template, args.output, data)

    print(f"✅ 已生成：{args.output}")
