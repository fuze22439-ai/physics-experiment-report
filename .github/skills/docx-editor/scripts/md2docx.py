"""
Markdown 转 docx 简易工具
将 Markdown 文件转为格式化的 Word 文档

用法:
    python md2docx.py input.md output.docx
"""
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx(md_path, docx_path):
    """将 Markdown 文件转换为 docx"""
    if not os.path.exists(md_path):
        print(f"错误：文件不存在 {md_path}")
        return False

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    lines = content.split('\n')
    in_table = False

    for line in lines:
        stripped = line.strip()

        # 空行
        if not stripped:
            if in_table:
                in_table = False
            continue

        # 标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = '黑体'

        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = '黑体'

        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = '黑体'

        # 表格
        elif stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
            # 表格处理略（简单的表格可以用逗号/制表符分隔格式）

        # 无序列表
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')

        # 有序列表
        elif re.match(r'^\d+\.\s', stripped):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')

        # 公式（保留 LaTeX 源码）
        elif stripped.startswith('$$') or stripped.startswith('$'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # 普通段落
        else:
            p = doc.add_paragraph(stripped)

    doc.save(docx_path)
    print(f"✅ Markdown → docx 完成：{docx_path}")
    return True


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法：python md2docx.py input.md output.docx")
        sys.exit(1)

    md_to_docx(sys.argv[1], sys.argv[2])
