"""
基于实验报告模板生成 .docx 报告

依赖：pip install python-docx

用法：
    python gen_report.py <实验文件夹名> [--name 姓名] [--student-id 学号]
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


TEMPLATE_PATH = r'E:\物理实验报告II\实验报告模版.docx'


def fill_report(template_path, output_path, data: dict):
    """
    基于模板填充实验报告

    Parameters
    ----------
    template_path : str
        模板 .docx 文件路径
    output_path : str
        输出 .docx 文件路径
    data : dict
        包含以下键的报告内容：
        - title: 实验项目名称
        - name: 姓名
        - student_id: 学号
        - date: 日期 (月/日)
        - purpose: 实验目的
        - equipment: 实验仪器与用具
        - principle: 实验原理
        - procedure_and_data: 实验要求及数据记录
        - analysis: 数据处理及图形
        - discussion: 结果分析及讨论
    """
    doc = Document(template_path)

    # 遍历所有段落，查找并替换占位符
    # 模板结构：基于 MarkItDown 转换结果，定位各部分

    paragraphs = doc.paragraphs

    # 找到关键段落索引
    section_markers = {
        'title': '实验项目：',
        'purpose': '【实验目的】',
        'equipment': '【实验仪器与用具】',
        'principle': '【实验原理】',
        'procedure': '【实验要求及数据记录】',
        'analysis': '【数据处理及图形】',
        'discussion': '【结果分析及讨论】',
    }

    section_indices = {}
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        for key, marker in section_markers.items():
            if marker in text:
                section_indices[key] = i

    # 填充标题行
    if 'title' in section_indices:
        idx = section_indices['title']
        paragraphs[idx].text = f'实验项目：{data.get("title", "")}'

    # 填充姓名学号日期行
    if 'title' in section_indices:
        name_line_idx = section_indices['title'] + 1
        if name_line_idx < len(paragraphs):
            name = data.get('name', '')
            sid = data.get('student_id', '')
            date_str = data.get('date', f'{datetime.now().month}月{datetime.now().day}日')
            paragraphs[name_line_idx].text = f'姓 名 {name} 学 号 {sid} 日期 {date_str}'

    def set_section_content(marker_key, content):
        """在对应标记段落后插入内容"""
        if marker_key not in section_indices:
            return
        idx = section_indices[marker_key]
        # 处理多行内容
        lines = content.strip().split('\n')
        # 第一行替换标记段落
        if lines:
            paragraphs[idx].text = section_markers[marker_key]
            # 在标记段落后插入内容段落
            for line in lines:
                if line.strip():
                    # 在当前段落后插入新段落
                    # 注意：python-docx 的 paragraph 插入比较复杂
                    pass  # 实际实现需要更精细的段落操作

    # 简单实现：直接替换文本
    # 由于 docx 模板格式复杂，建议使用更可靠的方法

    # 遍历所有段落进行文本替换
    for para in paragraphs:
        full_text = para.text

        # 替换实验项目
        if '实验项目：' in full_text:
            # 保留原有格式，仅修改文本
            for run in para.runs:
                if '实验项目：' in run.text:
                    run.text = f'实验项目：{data.get("title", "")}'
                    break

        # 替换姓名学号日期
        if '姓 名' in full_text and '学 号' in full_text:
            for run in para.runs:
                if '姓 名' in run.text:
                    name = data.get('name', '')
                    sid = data.get('student_id', '')
                    date_str = data.get('date', f'{datetime.now().month}月{datetime.now().day}日')
                    run.text = f'姓 名 {name} 学 号 {sid} 日期 {date_str}'
                    break

    # 更可靠的方法：通过 add_paragraph 在标记后插入内容
    # 先清除标记段落后的占位内容，再插入实际内容

    def fill_section_after_marker(marker, content_text):
        """在标记段落后填充内容"""
        for i, para in enumerate(paragraphs):
            if marker in para.text:
                # 找到标记段落，在其后插入内容
                # 如果标记段落后的段落是空的占位符，替换之
                if i + 1 < len(paragraphs):
                    next_para = paragraphs[i + 1]
                    # 检查是否是占位符（括号开头等）
                    next_text = next_para.text.strip()
                    if next_text.startswith('(') or next_text == '':
                        # 替换占位符
                        next_para.text = ''
                        for j, line in enumerate(content_text.strip().split('\n')):
                            if j == 0:
                                next_para.text = line
                            else:
                                # 在 next_para 后插入新段落
                                new_para = doc.add_paragraph()
                                # 复制样式
                                new_para.style = next_para.style
                                new_para.text = line
                                # 将新段落移动到正确位置
                                # python-docx 不支持直接移动段落，需要其他方法
                        break
                break

    # 填充各部分内容
    sections_content = {
        '【实验目的】': data.get('purpose', ''),
        '【实验仪器与用具】': data.get('equipment', ''),
        '【实验原理】': data.get('principle', ''),
        '【实验要求及数据记录】': data.get('procedure_and_data', ''),
        '【数据处理及图形】': data.get('analysis', ''),
        '【结果分析及讨论】': data.get('discussion', ''),
    }

    for marker, content in sections_content.items():
        if content:
            fill_section_after_marker(marker, content)

    doc.save(output_path)
    print(f'✅ 报告已生成：{output_path}')


def generate_from_markdown(md_content: str, output_path: str):
    """
    从 Markdown 内容生成 .docx 报告（替代方案，不依赖模板）

    使用 python-docx 从头创建，但格式需要手动设置
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 解析 Markdown 并生成
    lines = md_content.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(18)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(14)
        elif line.startswith('|'):
            # 表格处理略（需要更复杂的解析）
            p = doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f'✅ 报告已生成：{output_path}')


if __name__ == '__main__':
    # 示例用法
    sample_data = {
        'title': '扫描光电流实验',
        'name': '张三',
        'student_id': '2024000001',
        'date': '5月12日',
        'purpose': '1. 掌握扫描光电流测量方法\n2. 测量半导体材料的扩散长度',
        'equipment': '- 激光器：波长 532nm\n- 三维平移台：精度 1μm\n- 锁相放大器：SR830',
        'principle': '当激光照射半导体表面时，产生非平衡载流子...',
        'procedure_and_data': '1. 搭建光路\n2. 进行线扫描\n\n| 电压(V) | 峰值电流(μA) | 峰值位置(μm) |\n|---------|-------------|-------------|',
        'analysis': '采用指数衰减公式 $I = I_0 e^{-|x-D|/L}$ 拟合...',
        'discussion': '实验测得扩散长度 L = 667 ± 12 μm...',
    }

    output = r'E:\物理实验报告II\扫描光电流\实验报告_生成.docx'
    fill_report(TEMPLATE_PATH, output, sample_data)
