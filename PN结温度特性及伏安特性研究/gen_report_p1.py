"""
生成 PN 结温度特性及伏安特性研究 实验报告（前三部分）
基于模板填充：实验目的、实验仪器与用具、实验原理
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import copy, re

TEMPLATE = r"E:\物理实验报告II\PN结温度特性及伏安特性研究\实验报告.docx"

doc = Document(TEMPLATE)

def set_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)

def clear_runs(para):
    for r in para.runs:
        r.text = ""

def write_para(para, text):
    clear_runs(para)
    if para.runs:
        para.runs[0].text = text
        set_black(para.runs[0])
    else:
        r = para.add_run(text)
        set_black(r)

# ---- 基本信息 ----
for p in doc.paragraphs:
    txt = p.text
    if "实验项目" in txt:
        write_para(p, "实验项目：PN结温度特性及伏安特性研究")
    # 姓名学号行：模板特征是含"姓"和"学号"和"日期"
    if ("姓" in txt or "名" in txt) and ("学号" in txt or "学 号" in txt):
        write_para(p, "姓名______________ 学号______________ 日期____月____日")

# ---- 【实验目的】 ----
purpose = (
    "1. 研究PN结正向电压随温度变化的基本规律，理解PN结温度传感器的物理机制；\n"
    "2. 测量并绘制PN结正向伏安特性曲线，验证PN结的指数型电流-电压关系；\n"
    "3. 在恒定正向电流条件下，测绘PN结正向压降随温度变化的关系曲线；\n"
    "4. 测定PN结温度传感器的灵敏度S，并估算被测PN结材料的禁带宽度Eg(0)。"
)

# ---- 【实验仪器与用具】 ----
instruments = (
    "PN结正向特性综合实验仪（含电流源、电压表）；"
    "温度传感实验装置（含加热电流源、温控系统）；"
    "样品室（内嵌加热模块）；"
    "Pt100铂电阻温度传感器（测温精度 ±0.1°C）；"
    "PN结集成温度传感器（硅三极管共基极接法）；"
    "计算机及虚拟仿真软件"
)

# ---- 【实验原理】 ----
principle = (
    "1. PN结的伏安特性\n"
    "根据半导体物理理论，理想PN结的正向电流IF与正向压降VF满足以下关系：\n"
    "IF = IS [exp(qVF/kT) − 1]    (1)\n"
    "其中IS为反向饱和电流，q为电子电荷量（1.602×10⁻¹⁹ C），k为玻耳兹曼常数，"
    "T为热力学温度。在常温下，exp(qVF/kT) ≫ 1，上式可简化为：\n"
    "IF = IS exp(qVF/kT)    (2)\n"
    "此即PN结正向伏安特性的基本方程，表明在恒定温度下IF与VF呈指数关系。\n\n"
    "2. PN结温度传感器的基本原理\n"
    "反向饱和电流IS与PN结材料的禁带宽度及温度密切相关，其表达式为：\n"
    "IS = C·Tʳ·exp[qVg(0)/kT]    (3)\n"
    "式中C为与结面积、掺杂浓度相关的常数，r为常数，Vg(0)为绝对零度时导带底"
    "与价带顶的电势差。将式(3)代入式(2)，两边取对数整理得：\n"
    "VF = Vg(0) − (k/q·ln(C/IF))·T − (kT/q)·ln(Tʳ) = V₁ + Vₙ₁    (4)\n"
    "其中V₁为线性项，Vₙ₁为非线性项。对于硅PN结，在−50~150°C范围内Vₙ₁可忽略，"
    "VF随T近似呈线性下降。此即PN结温度传感器的基本工作原理。\n\n"
    "3. 禁带宽度的测量\n"
    "忽略非线性项后，式(4)简化为：\n"
    "VF = Vg(0) − (k/q·ln(C/IF))·T = Vg(0) + S·T    (5)\n"
    "式中S = ΔVF/ΔT（单位mV/°C）即为PN结温度传感器的灵敏度。"
    "在恒定电流IF下，通过实验测量VF−T关系曲线，由斜率求得S。"
    "进而可由Vg(0) = VF − S·T求出绝对零度时的电势差，"
    "禁带宽度Eg(0) = q·Vg(0)。硅的Eg(0)理论值约为1.21 eV。\n\n"
    "4. 玻耳兹曼常数的测量\n"
    "由式(2)，在恒定温度T下，取两组(IF, VF)数据，可求得玻耳兹曼常数：\n"
    "k = (q/T)·(VF1 − VF2) / ln(IF2/IF1)    (6)\n"
    "为提高精度，通常采用指数函数曲线回归法：令IF = A·exp(B·VF)，"
    "其中A = IS，B = q/kT，通过多组(IF, VF)数据的指数回归拟合求得B值，"
    "进而得到k = q/(B·T)。\n\n"
    "※ 实验原理示意图请参见书本图5.13.1～图5.13.5，请在报告中插入对应图片。"
)

# ---- 填写三个板块 ----
section_map = {
    "【实验目的】": purpose,
    "【实验仪器与用具】": instruments,
    "【实验原理】": principle,
}

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt in section_map:
        # 标题段：确保黑色
        for r in p.runs:
            set_black(r)
        # 下一个段落是占位内容，填入正文
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            write_para(next_p, section_map[txt])

# ---- 全局黑色 ----
for p in doc.paragraphs:
    for r in p.runs:
        set_black(r)

doc.save(TEMPLATE)
print("✅ 报告前三部分已生成：", TEMPLATE)
