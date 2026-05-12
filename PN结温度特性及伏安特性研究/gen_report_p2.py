"""
生成 PN 结温度特性及伏安特性研究 实验报告（前四部分）
基于模板填充：实验目的、实验仪器与用具、实验原理、实验要求及数据记录
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
import copy, re

TEMPLATE = r"E:\物理实验报告II\PN结温度特性及伏安特性研究\实验报告.docx"

doc = Document(TEMPLATE)

def set_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)

def clear_runs(para):
    for r in para.runs:
        r.text = ""

def write_para(para, text):
    clear_runs(para)
    if para.runs:
        para.runs[0].text = text
        set_black(para.runs[0])
    else:
        r = para.add_run(text)
        set_black(r)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def make_table_cell_text(cell, text, bold=False, font_size=Pt(9)):
    """设置单元格文本"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = 1  # 居中
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.name = "宋体"
    run.bold = bold
    set_black(run)
    # 设置段落间距为0
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

# ============================================================
# 基本信息
# ============================================================
for p in doc.paragraphs:
    txt = p.text
    if "实验项目" in txt:
        write_para(p, "实验项目：PN结温度特性及伏安特性研究")
    if ("姓" in txt or "名" in txt) and ("学号" in txt or "学 号" in txt):
        write_para(p, "姓名______________ 学号______________ 日期____月____日")

# ============================================================
# 各板块内容
# ============================================================
purpose = (
    "1. 研究PN结正向电压随温度变化的基本规律，理解PN结温度传感器的物理机制；\n"
    "2. 测量并绘制PN结正向伏安特性曲线，验证PN结的指数型电流-电压关系；\n"
    "3. 在恒定正向电流条件下，测绘PN结正向压降随温度变化的关系曲线；\n"
    "4. 测定PN结温度传感器的灵敏度S，并估算被测PN结材料的禁带宽度Eg(0)。"
)

instruments = (
    "PN结正向特性综合实验仪（含电流源、电压表）；"
    "温度传感实验装置（含加热电流源、温控系统）；"
    "样品室（内嵌加热模块）；"
    "Pt100铂电阻温度传感器（测温精度 ±0.1°C）；"
    "PN结集成温度传感器（硅三极管共基极接法）；"
    "计算机及虚拟仿真软件"
)

principle = (
    "1. PN结的伏安特性\n"
    "根据半导体物理理论，理想PN结的正向电流IF与正向压降VF满足以下关系：\n"
    "I_F = I_S \\left[ \\exp \\left( \\frac{q V_F}{k T} \\right) - 1 \\right]    (5.13.1)\n"
    "其中IS为反向饱和电流（在温度恒定时为常数），q为电子电荷量（1.602×10⁻¹⁹ C），"
    "k为玻耳兹曼常数，T为热力学温度。在常温下，exp(qVF/kT) ≫ 1，上式可简化为：\n"
    "I_F = I_S \\exp \\left( \\frac{q V_F}{k T} \\right)    (5.13.2)\n"
    "此即PN结正向伏安特性的基本方程，表明在恒定温度下IF与VF呈指数关系。\n\n"
    "2. PN结温度传感器的基本原理\n"
    "反向饱和电流IS与PN结材料的禁带宽度及温度密切相关，其表达式为：\n"
    "I_S = C T^r \\exp \\left[ \\frac{q V_g (0)}{k T} \\right]    (5.13.3)\n"
    "式中C为与PN结面积、掺杂浓度等有关的常数，r也是常数，Vg(0)为绝对零度时"
    "PN结材料的导带底与价带顶的电势差。\n"
    "将式(5.13.3)代入式(5.13.2)，两边取对数整理得：\n"
    "V_F = V_g (0) - \\left( \\frac{k}{q} \\ln \\frac{C}{I_F} \\right) T"
    " - \\frac{k T}{q} \\ln T^r = V_1 + V_{n1}    (5.13.4)\n"
    "其中：\n"
    "V_1 = V_g (0) - \\left( \\frac{k}{q} \\ln \\frac{C}{I_F} \\right) T,"
    " \\quad V_{n1} = - \\frac{k T}{q} \\ln T^r    (5.13.5)\n"
    "式(5.13.4)是PN结正向压降作为电流和温度函数的表达式，也是PN结温度传感器的"
    "基本方程。当IF为常数时，正向压降VF只随温度变化。V₁为线性项，Vₙ₁为非线性项。"
    "对于通常的硅PN结材料，在温度范围−50~150°C时，Vₙ₁的影响可以忽略不计。"
    "在恒定小电流条件下，PN结正向压降随温度升高几乎呈线性下降。\n\n"
    "3. 禁带宽度的测量\n"
    "忽略非线性项Vₙ₁后，由式(5.13.4)和(5.13.5)可得PN结正向压降VF与热力学温度T"
    "关系的近似关系式：\n"
    "V_F = V_g (0) - \\left( \\frac{k}{q} \\ln \\frac{C}{I_F} \\right) T"
    " = V_g (0) + S T    (5.13.6)\n"
    "式中S = ΔVF/ΔT（单位mV/°C）即为PN结温度传感器的灵敏度。"
    "在恒定电流IF下，通过实验测量VF−T关系曲线，由斜率求得S。\n"
    "根据式(5.13.6)可得：\n"
    "V_g (0) = V_F - S T    (5.13.7)\n"
    "进而求出绝对零度时半导体材料的近似禁带宽度：\n"
    "E_g (0) = q V_g (0)\n"
    "硅材料的Eg(0)约为1.21 eV。\n\n"
    "4. 玻耳兹曼常数的测量\n"
    "由式(5.13.2)，在恒定温度T下，取两组(IF, VF)数据，可求得玻耳兹曼常数：\n"
    "k = \\frac{q}{T} \\cdot \\frac{V_{F1} - V_{F2}}{\\ln (I_{F2} / I_{F1})}    (5.13.8)\n"
    "为提高精度，通常采用指数函数曲线回归法。将式(5.13.2)进行变量代换：\n"
    "I_F = A \\exp (B V_F)\n"
    "其中A = IS，B = q/kT。以IF和VF为变量，根据测得的数据进行指数函数的曲线回归，"
    "求得A、B值，进而求出反向饱和电流和玻耳兹曼常数k。\n\n"
    "※ 实验原理示意图请参见书本图5.13.1～图5.13.5，请在报告中插入对应图片。"
)

steps = (
    "（一）主窗口介绍\n"
    "成功进入实验场景后，实验场景的主窗口如图5.13.6所示。\n"
    "[图片：图5.13.6 PN结物理特性测试实验主场景图]\n\n"
    "（二）放置PN结和Pt100温度传感器到样品室中\n"
    "用鼠标依次拖动Pt100电阻和PN结放置到样品室上，当鼠标松开时，仪器被插入"
    "到样品室插孔中，如图5.13.7所示。\n"
    "[图片：图5.13.7 放置PN结和Pt100温度传感器到样品室中]\n\n"
    "（三）实验连线\n"
    "当鼠标移动到实验仪器接线柱的上方，拖动鼠标，便会产生\"导线\"，当鼠标移动到"
    "另一个接线柱的时候，松开鼠标，两个接线柱之间便产生一条导线，连线成功；"
    "如果松开鼠标的时候，鼠标不是在某个接线柱上，画出的导线将会被自动销毁，"
    "此次连线失败，如图5.13.8所示。根据电路图连接好电路，然后在数据表格中"
    "点击\"连线\"模块下的\"确定状态\"按钮，保存连线状态。\n"
    "[图片：图5.13.8 实验连线示意图]\n\n"
    "（四）测量PN结正向伏安特性 IF-VF 曲线\n"
    "1. 设置实验温度：在主场景中双击打开温度传感实验装置；在仪器背面视图中"
    "点击电源开关，打开电源；点击\"SET\"按钮，设置温度值为30°C。打开加热电流"
    "开关，并选择合适的加热电流给样品室进行加热；待温度恒定后开始实验。\n"
    "2. 绘制30°C时的PN结正向伏安特性IF-VF曲线：在主场景中双击打开PN结正向"
    "特性综合实验仪，在仪器背面点击电源开关，打开电源；选择合适的电流量程"
    "挡位，并调节电流旋钮，使IF逐渐增大，正向压降VF将随之增大。要求VF在"
    "0.450~0.540 V范围内每变化0.005 V记录对应的IF，数据记录于表5.13.1。\n"
    "[图片：图5.13.9 温度设置示意图] [图片：图5.13.10 IF-VF曲线实验装置]\n\n"
    "（五）测绘PN结正向压降VF随温度的变化曲线\n"
    "1. 调节电流IF = 50 μA：在主场景中双击打开PN结正向特性综合实验仪；在仪器"
    "背面点击电源开关，打开电源；选择合适的电流量程挡位，并调节电流旋钮，"
    "使电流IF等于50 μA（正向电流IF一般选小于100 μA，不宜太大）。\n"
    "[图片：图5.13.11 调节电流IF = 50 μA的实验装置]\n"
    "2. 在恒定电流IF = 50 μA条件下，测绘PN结正向压降VF随温度的变化曲线。"
    "要求在30~80°C温度范围（温度不宜太高）内每隔5°C测量一个点，升温过程"
    "和降温过程各测一遍，数据记录于表5.13.2。\n"
    "[图片：图5.13.12 相关装置]\n\n"
    "3. 计算玻耳兹曼常数k：根据表5.13.1的测量结果，依据指数函数的曲线回归"
    "计算玻耳兹曼常数k，并与公认值k = 1.38×10⁻²³ J/K比较，计算误差。\n"
    "4. 计算灵敏度S和禁带宽度Eg(0)：根据表5.13.2测得的数据，由公式(5.13.6)"
    "求被测PN结正向压降随温度变化的灵敏度S（mV/°C）；并根据式(5.13.7)，"
    "估算被测PN结材料的禁带宽度Eg(0)。"
)

# ============================================================
# 填写四个板块（前三个：段落替换）
# ============================================================
section_map = {
    "【实验目的】": purpose,
    "【实验仪器与用具】": instruments,
    "【实验原理】": principle,
}

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt in section_map:
        for r in p.runs:
            set_black(r)
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            write_para(next_p, section_map[txt])

# ============================================================
# 【实验要求及数据记录】 — 需要插入表格，特殊处理
# ============================================================
# 找到该段落及其后续占位段
step_section_idx = None
step_content_idx = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if "【实验要求及数据记录】" in txt:
        for r in p.runs:
            set_black(r)
        step_section_idx = i
        if i + 1 < len(doc.paragraphs):
            step_content_idx = i + 1
        break

if step_content_idx is not None:
    step_para = doc.paragraphs[step_content_idx]
    write_para(step_para, steps)

# ============================================================
# 在步骤段落后插入两个数据表格
# ============================================================

def insert_element_after(anchor, new_element):
    """在指定段落或XML元素后插入 XML 元素"""
    if hasattr(anchor, '_element'):
        anchor._element.addnext(new_element)
    else:
        # anchor 本身就是 XML 元素
        anchor.addnext(new_element)

def make_paragraph_element(text, bold=False, font_size=Pt(10)):
    """创建一个新段落的 XML 元素"""
    p_elem = OxmlElement('w:p')
    r_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    r_elem.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')
    r_elem.append(t_elem)
    p_elem.append(r_elem)
    return p_elem

# --- 表5.13.1 PN结正向伏安特性 IF-VF 曲线 ---
# 原始格式：4行 × 10列（分两半显示 9+9 个 VF 值）
vf_row1 = [0.450, 0.455, 0.460, 0.465, 0.470, 0.475, 0.480, 0.485, 0.490]
vf_row2 = [0.495, 0.500, 0.505, 0.510, 0.515, 0.520, 0.525, 0.530, 0.535]

# 在步骤内容段落后插入表标题
insert_element_after(doc.paragraphs[step_content_idx],
                     make_paragraph_element("表5.13.1  PN结正向伏安特性 IF-VF 曲线  （T = 30°C）",
                                           bold=True, font_size=Pt(10)))

# 创建表5.13.1：4行 × 10列（第一列为标签 + 9个数据列）
num_cols = 10
table1 = doc.add_table(rows=4, cols=num_cols)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# 行0：VF/V（上半）
make_table_cell_text(table1.cell(0, 0), "VF / V", bold=True)
for j, vf in enumerate(vf_row1):
    make_table_cell_text(table1.cell(0, j + 1), f"{vf:.3f}", bold=True)

# 行1：IF/μA（上半数据，空）
make_table_cell_text(table1.cell(1, 0), "IF / μA", bold=True)
for j in range(len(vf_row1)):
    make_table_cell_text(table1.cell(1, j + 1), "")

# 行2：VF/V（下半）
make_table_cell_text(table1.cell(2, 0), "VF / V", bold=True)
for j, vf in enumerate(vf_row2):
    make_table_cell_text(table1.cell(2, j + 1), f"{vf:.3f}", bold=True)

# 行3：IF/μA（下半数据，空）
make_table_cell_text(table1.cell(3, 0), "IF / μA", bold=True)
for j in range(len(vf_row2)):
    make_table_cell_text(table1.cell(3, j + 1), "")

# 把 table1 移动到步骤段落后
tbl1_elem = table1._tbl
doc.element.body.remove(tbl1_elem)
# 找步骤段落后第二个段落（即刚插入的标题段）
step_elem = doc.paragraphs[step_content_idx]._element
title_elem = step_elem.getnext()  # 刚插入的标题段落
if title_elem is not None:
    title_elem.addnext(tbl1_elem)
else:
    doc.element.body.append(tbl1_elem)

# --- 表5.13.2 ---
temps = [30.0, 35.0, 40.0, 45.0, 50.0, 55.0, 60.0, 65.0, 70.0, 75.0]

# 在表1后插入空行 + 表标题
tbl1_elem = table1._tbl
insert_element_after(tbl1_elem,
                     make_paragraph_element("", bold=False, font_size=Pt(6)))
# 找到刚插入的空段落元素
empty_elem = tbl1_elem.getnext()
insert_element_after(empty_elem,
                     make_paragraph_element("表5.13.2  PN结正向压降 VF 随温度的变化曲线  （IF = 50 μA）",
                                           bold=True, font_size=Pt(10)))

# 创建表5.13.2
table2 = doc.add_table(rows=11, cols=4)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ["温度值 / °C", "升温 VF / V", "降温 VF / V", "VF 平均值 / V"]
for j, h in enumerate(headers2):
    make_table_cell_text(table2.cell(0, j), h, bold=True)

for i, t in enumerate(temps):
    make_table_cell_text(table2.cell(i + 1, 0), f"{t:.1f}")
    make_table_cell_text(table2.cell(i + 1, 1), "")
    make_table_cell_text(table2.cell(i + 1, 2), "")
    make_table_cell_text(table2.cell(i + 1, 3), "")

# 移动 table2 到标题段落后
doc.element.body.remove(table2._tbl)
title2_elem = empty_elem.getnext()  # 这是表5.13.2的标题段落
if title2_elem is not None:
    title2_elem.addnext(table2._tbl)
else:
    doc.element.body.append(table2._tbl)

# ============================================================
# 全局黑色字体 + 清理残留占位符
# ============================================================
for p in doc.paragraphs:
    for r in p.runs:
        set_black(r)
    # 清理模板残留占位符
    if "（给出详细数据记录表格）" in p.text:
        write_para(p, "")
    if "（简要说明操作步骤）" in p.text:
        write_para(p, "")

doc.save(TEMPLATE)
print("✅ 报告前四部分已生成：", TEMPLATE)
